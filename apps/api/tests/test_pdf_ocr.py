"""扫描型/混合型 PDF → TXT/DOCX 的 OCR 回归测试。"""
from __future__ import annotations

import asyncio
import threading
from pathlib import Path
from types import SimpleNamespace

import pytest
from PIL import Image, ImageDraw
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

import app.handlers.pdf_handlers as pdf_handlers
import app.providers as providers
from app.errors import ConversionError, ErrorCode
from app.handlers.pdf_handlers import PDF_HANDLERS
from app.providers.base import OCRProvider, OCRResult


class RecordingOCR(OCRProvider):
    def __init__(self, text: str = "扫描页 OCR 文字") -> None:
        self.text = text
        self.calls: list[dict] = []

    async def extract_text(
        self,
        image_path: str,
        language: str = "auto",
        detect_tables: bool = True,
        preserve_layout: bool = False,
    ) -> OCRResult:
        path = Path(image_path)
        assert path.suffix == ".png"
        assert path.exists()
        self.calls.append(
            {
                "language": language,
                "detect_tables": detect_tables,
                "preserve_layout": preserve_layout,
            }
        )
        return OCRResult(text=self.text, tables=[], confidence=0.9)


class FailingOCR(OCRProvider):
    async def extract_text(
        self,
        image_path: str,
        language: str = "auto",
        detect_tables: bool = True,
        preserve_layout: bool = False,
    ) -> OCRResult:
        raise ConversionError(ErrorCode.OCR_FAILED, "扫描页 OCR 失败")


class CoordinatedOCR(OCRProvider):
    def __init__(
        self,
        delays: dict[int, float] | None = None,
        fail_pages: set[int] | None = None,
        wait_for_started: int | None = None,
    ) -> None:
        self.delays = delays or {}
        self.fail_pages = fail_pages or set()
        self.wait_for_started = wait_for_started
        self._started_barrier = asyncio.Event()
        self.started: list[int] = []
        self.active = 0
        self.max_active = 0

    async def extract_text(
        self,
        image_path: str,
        language: str = "auto",
        detect_tables: bool = True,
        preserve_layout: bool = False,
    ) -> OCRResult:
        page_number = int(Path(image_path).stem.rsplit("-", 1)[1])
        self.started.append(page_number)
        self.active += 1
        self.max_active = max(self.max_active, self.active)
        try:
            if (
                self.wait_for_started is not None
                and len(self.started) >= self.wait_for_started
            ):
                self._started_barrier.set()
            if self.wait_for_started is not None:
                await self._started_barrier.wait()
            await asyncio.sleep(self.delays.get(page_number, 0.03))
            if page_number in self.fail_pages:
                raise ConversionError(ErrorCode.OCR_FAILED, f"第 {page_number} 页 OCR 失败")
            return OCRResult(
                text=f"OCR PAGE {page_number}", tables=[], confidence=0.9,
            )
        finally:
            self.active -= 1


def _fake_render_pdf_page(
    input_path: str, page_index: int, dest: Path,
) -> None:
    dest.write_bytes(f"page-{page_index + 1}".encode())


def _ocr_states(count: int) -> list[pdf_handlers.PdfPageState]:
    return [
        pdf_handlers.PdfPageState(index=index, native_text="", needs_ocr=True)
        for index in range(count)
    ]


@pytest.fixture
def ocr_provider(monkeypatch) -> RecordingOCR:
    provider = RecordingOCR()
    monkeypatch.setattr(providers, "_ocr", provider)
    return provider


def _run(
    key: str,
    input_path: Path,
    output_dir: Path,
    target_ext: str,
    options: dict,
) -> dict:
    return asyncio.run(
        PDF_HANDLERS[key].run(
            str(input_path), str(output_dir), target_ext, options,
        )
    )


def _make_scan_image(path: Path) -> Path:
    image = Image.new("RGB", (1200, 800), "white")
    ImageDraw.Draw(image).text((120, 320), "SCANNED PAGE 123", fill="black")
    image.save(path, "PNG")
    return path


def _make_scanned_pdf(path: Path) -> Path:
    image_path = _make_scan_image(path.with_suffix(".png"))
    with Image.open(image_path) as image:
        image.convert("RGB").save(path, "PDF", resolution=150)
    return path


def _make_multi_page_scanned_pdf(path: Path, page_count: int = 3) -> Path:
    scan_image = _make_scan_image(path.with_suffix(".png"))
    width, height = A4
    pdf = canvas.Canvas(str(path), pagesize=A4)
    for _ in range(page_count):
        pdf.drawImage(ImageReader(str(scan_image)), 0, 0, width=width, height=height)
        pdf.showPage()
    pdf.save()
    return path


def _make_native_pdf(path: Path) -> Path:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.drawString(72, 760, "NATIVE PAGE TEXT")
    pdf.save()
    return path


def _make_mixed_pdf(path: Path) -> Path:
    scan_image = _make_scan_image(path.with_suffix(".png"))
    width, height = A4
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.drawString(72, 760, "NATIVE PAGE TEXT")
    pdf.showPage()
    pdf.drawImage(ImageReader(str(scan_image)), 0, 0, width=width, height=height)
    pdf.save()
    return path


def test_scanned_pdf_to_txt_uses_ocr_and_forwards_options(
    tmp_path: Path, ocr_provider: RecordingOCR,
) -> None:
    src = _make_scanned_pdf(tmp_path / "scan.pdf")
    out_dir = tmp_path / "out"
    out_dir.mkdir()

    result = _run(
        "pdf_to_txt",
        src,
        out_dir,
        ".txt",
        {"ocr_language": "zh", "detect_tables": False, "preserve_layout": True},
    )

    assert "扫描页 OCR 文字" in Path(result["output_path"]).read_text(encoding="utf-8")
    assert ocr_provider.calls == [
        {"language": "zh", "detect_tables": False, "preserve_layout": True}
    ]
    assert "OCR" in (result["quality_notice"] or "")
    assert not list(out_dir.glob("ocr-page-*.png"))


def test_scanned_pdf_to_docx_contains_editable_ocr_text(
    tmp_path: Path, ocr_provider: RecordingOCR,
) -> None:
    src = _make_scanned_pdf(tmp_path / "scan.pdf")
    out_dir = tmp_path / "out"
    out_dir.mkdir()

    result = _run("pdf_to_docx", src, out_dir, ".docx", {})

    text = "\n".join(p.text for p in Document(result["output_path"]).paragraphs)
    assert "扫描页 OCR 文字" in text
    assert len(ocr_provider.calls) == 1
    assert "OCR" in (result["quality_notice"] or "")


def test_mixed_pdf_keeps_page_order_and_ocrs_only_scan_page(
    tmp_path: Path, ocr_provider: RecordingOCR,
) -> None:
    src = _make_mixed_pdf(tmp_path / "mixed.pdf")
    out_dir = tmp_path / "out"
    out_dir.mkdir()

    result = _run("pdf_to_txt", src, out_dir, ".txt", {})

    text = Path(result["output_path"]).read_text(encoding="utf-8")
    assert text.index("NATIVE PAGE TEXT") < text.index("扫描页 OCR 文字")
    assert len(ocr_provider.calls) == 1


def test_pdf_ocr_failure_is_preserved_and_page_image_is_removed(
    tmp_path: Path, monkeypatch,
) -> None:
    monkeypatch.setattr(providers, "_ocr", FailingOCR())
    src = _make_scanned_pdf(tmp_path / "scan.pdf")
    out_dir = tmp_path / "out"
    out_dir.mkdir()

    with pytest.raises(ConversionError) as exc:
        _run("pdf_to_txt", src, out_dir, ".txt", {})

    assert exc.value.code == ErrorCode.OCR_FAILED
    assert not list(out_dir.glob("ocr-page-*.png"))


def test_native_pdf_docx_does_not_call_ocr(
    tmp_path: Path, ocr_provider: RecordingOCR,
) -> None:
    src = _make_native_pdf(tmp_path / "native.pdf")
    out_dir = tmp_path / "out"
    out_dir.mkdir()

    result = _run("pdf_to_docx", src, out_dir, ".docx", {})

    Document(result["output_path"])
    assert len(ocr_provider.calls) == 0


@pytest.mark.parametrize(("limit", "expected_peak"), [(1, 1), (2, 2), (3, 3)])
def test_pdf_ocr_respects_page_concurrency_and_keeps_order(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    limit: int,
    expected_peak: int,
) -> None:
    provider = CoordinatedOCR(
        delays={1: 0.08, 2: 0.06, 3: 0.04, 4: 0.02, 5: 0.01},
    )
    monkeypatch.setattr(providers, "_ocr", provider)
    monkeypatch.setattr(pdf_handlers, "_render_pdf_page", _fake_render_pdf_page)
    monkeypatch.setattr(
        pdf_handlers,
        "settings",
        SimpleNamespace(pdf_ocr_page_concurrency=limit),
    )

    result = asyncio.run(
        pdf_handlers._extract_pdf_text_with_ocr(
            "unused.pdf", tmp_path, {}, _ocr_states(5),
        )
    )

    assert provider.max_active == expected_peak
    assert result.text.split("\n\n") == [f"OCR PAGE {n}" for n in range(1, 6)]
    assert result.ocr_pages == 5
    assert not list(tmp_path.glob("ocr-page-*.png"))


def test_pdf_ocr_failure_stops_new_pages_and_cleans_inflight_images(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    provider = CoordinatedOCR(
        delays={1: 0.08, 2: 0.01, 3: 0.08},
        fail_pages={2},
    )
    monkeypatch.setattr(providers, "_ocr", provider)
    monkeypatch.setattr(pdf_handlers, "_render_pdf_page", _fake_render_pdf_page)
    monkeypatch.setattr(
        pdf_handlers,
        "settings",
        SimpleNamespace(pdf_ocr_page_concurrency=3),
    )

    with pytest.raises(ConversionError) as exc:
        asyncio.run(
            pdf_handlers._extract_pdf_text_with_ocr(
                "unused.pdf", tmp_path, {}, _ocr_states(7),
            )
        )

    assert exc.value.code == ErrorCode.OCR_FAILED
    assert set(provider.started) == {1, 2, 3}
    assert not list(tmp_path.glob("ocr-page-*.png"))


def test_pdf_ocr_multiple_failures_raise_lowest_page_error(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    provider = CoordinatedOCR(
        delays={1: 0.05, 2: 0.0},
        fail_pages={1, 2},
        wait_for_started=2,
    )
    monkeypatch.setattr(providers, "_ocr", provider)
    monkeypatch.setattr(pdf_handlers, "_render_pdf_page", _fake_render_pdf_page)
    monkeypatch.setattr(
        pdf_handlers,
        "settings",
        SimpleNamespace(pdf_ocr_page_concurrency=2),
    )

    with pytest.raises(ConversionError) as exc:
        asyncio.run(
            pdf_handlers._extract_pdf_text_with_ocr(
                "unused.pdf", tmp_path, {}, _ocr_states(2),
            )
        )

    assert exc.value.code == ErrorCode.OCR_FAILED
    assert exc.value.message == "第 1 页 OCR 失败"


def test_pdf_ocr_cancellation_waits_for_render_thread_before_cleanup(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    provider = RecordingOCR()
    render_started = threading.Event()
    allow_render_finish = threading.Event()
    render_finished = threading.Event()

    def blocking_render(input_path: str, page_index: int, dest: Path) -> None:
        dest.write_bytes(b"render-started")
        render_started.set()
        assert allow_render_finish.wait(timeout=2)
        dest.write_bytes(b"render-finished")
        render_finished.set()

    monkeypatch.setattr(providers, "_ocr", provider)
    monkeypatch.setattr(pdf_handlers, "_render_pdf_page", blocking_render)
    monkeypatch.setattr(
        pdf_handlers,
        "settings",
        SimpleNamespace(pdf_ocr_page_concurrency=3),
    )

    async def cancel_during_render() -> bool:
        task = asyncio.create_task(
            pdf_handlers._extract_pdf_text_with_ocr(
                "unused.pdf", tmp_path, {}, _ocr_states(1),
            )
        )
        assert await asyncio.to_thread(render_started.wait, 2)
        task.cancel()
        await asyncio.sleep(0.02)
        waited_for_thread = not task.done()
        allow_render_finish.set()
        with pytest.raises(asyncio.CancelledError):
            await task
        assert await asyncio.to_thread(render_finished.wait, 2)
        return waited_for_thread

    assert asyncio.run(cancel_during_render())
    assert provider.calls == []
    assert not list(tmp_path.glob("ocr-page-*.png"))


def test_multi_page_scanned_pdf_to_docx_keeps_page_order(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    provider = CoordinatedOCR(delays={1: 0.08, 2: 0.04, 3: 0.01})
    monkeypatch.setattr(providers, "_ocr", provider)
    monkeypatch.setattr(
        pdf_handlers,
        "settings",
        SimpleNamespace(pdf_ocr_page_concurrency=3),
    )
    src = _make_multi_page_scanned_pdf(tmp_path / "multi-scan.pdf")
    out_dir = tmp_path / "out"
    out_dir.mkdir()

    result = _run("pdf_to_docx", src, out_dir, ".docx", {})

    text = "\n".join(p.text for p in Document(result["output_path"]).paragraphs)
    assert text.index("OCR PAGE 1") < text.index("OCR PAGE 2")
    assert text.index("OCR PAGE 2") < text.index("OCR PAGE 3")
    assert provider.max_active >= 2
    assert not list(out_dir.glob("ocr-page-*.png"))
