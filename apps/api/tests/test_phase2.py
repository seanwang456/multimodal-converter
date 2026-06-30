"""Phase 2 真实转换 handler 验证（生成样本，逐条核实）。office_to_pdf 需 libreoffice，本地未装则跳过。"""
from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

import pytest

from app.config import settings
from app.handlers.audio_handlers import AUDIO_HANDLERS
from app.handlers.image_handlers import IMAGE_HANDLERS
from app.handlers.office_handlers import OFFICE_HANDLERS
from app.handlers.txt_handlers import TXT_HANDLERS


def _run(handler, input_path: Path, out_dir: Path, ext: str) -> dict:
    import asyncio

    return asyncio.run(handler.run(str(input_path), str(out_dir), ext, {}))


# ---------- 样本生成 ----------
def _make_txt(p: Path) -> Path:
    p.write_text("标题行\n姓名,金额\n张三,100\n李四,200\n", encoding="utf-8")
    return p


def _make_png(p: Path) -> Path:
    from PIL import Image

    Image.new("RGB", (16, 16), (10, 20, 30)).save(str(p))
    return p


def _make_wav(p: Path) -> Path:
    subprocess.run(
        [settings.ffmpeg_bin, "-y", "-f", "lavfi", "-i", "sine=frequency=440:duration=1",
         "-ar", "22050", str(p)], check=True, capture_output=True
    )
    return p


def _make_docx(p: Path) -> Path:
    from docx import Document

    d = Document()
    d.add_paragraph("Hello 文档")
    d.save(str(p))
    return p


def _make_xlsx(p: Path) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["姓名", "金额"])
    ws.append(["张三", 100])
    wb.save(str(p))
    return p


# ---------- txt ----------
@pytest.mark.parametrize("target,openable", [(".docx", "docx"), (".pdf", "pdf"), (".pptx", "pptx"), (".xlsx", "xlsx")])
def test_txt_conversions(tmp_path, target, openable) -> None:
    src = _make_txt(tmp_path / "s.txt")
    out_dir = tmp_path / "out"
    out_dir.mkdir()
    key = {".docx": "txt_to_docx", ".pdf": "txt_to_pdf", ".pptx": "txt_to_pptx", ".xlsx": "txt_to_xlsx"}[target]
    res = _run(TXT_HANDLERS[key], src, out_dir, target)
    out = Path(res["output_path"])
    assert out.exists() and out.stat().st_size > 0
    assert out.suffix == target
    if openable == "docx":
        from docx import Document
        Document(str(out))
    elif openable == "pptx":
        from pptx import Presentation
        Presentation(str(out))
    elif openable == "xlsx":
        from openpyxl import load_workbook
        load_workbook(str(out))
    elif openable == "pdf":
        assert out.read_bytes()[:4] == b"%PDF"


# ---------- image ----------
def test_image_png_to_jpg(tmp_path) -> None:
    src = _make_png(tmp_path / "s.png")
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(IMAGE_HANDLERS["image_format_convert"], src, out_dir, ".jpg")
    out = Path(res["output_path"])
    from PIL import Image
    img = Image.open(out); img.load()
    assert out.suffix == ".jpg" and out.stat().st_size > 0


def test_image_jpg_to_png(tmp_path) -> None:
    src = tmp_path / "s.jpg"
    from PIL import Image
    Image.new("RGB", (12, 12), (1, 2, 3)).save(str(src), format="JPEG")
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(IMAGE_HANDLERS["image_format_convert"], src, out_dir, ".png")
    out = Path(res["output_path"])
    Image.open(out).verify()
    assert out.suffix == ".png" and out.stat().st_size > 0


def test_image_bmp_to_jpg(tmp_path) -> None:
    src = tmp_path / "s.bmp"
    from PIL import Image
    Image.new("RGB", (12, 12), (5, 6, 7)).save(str(src), format="BMP")
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(IMAGE_HANDLERS["image_format_convert"], src, out_dir, ".jpg")
    out = Path(res["output_path"])
    assert out.suffix == ".jpg" and out.stat().st_size > 0


# ---------- audio ----------
@pytest.mark.skipif(shutil.which(settings.ffmpeg_bin) is None, reason="ffmpeg 未安装")
@pytest.mark.parametrize("src_factory,src_ext,target", [
    (_make_wav, ".wav", ".mp3"),
])
def test_audio_wav_to_mp3(tmp_path, src_factory, src_ext, target) -> None:
    src = src_factory(tmp_path / ("s" + src_ext))
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(AUDIO_HANDLERS["audio_format_convert"], src, out_dir, target)
    out = Path(res["output_path"])
    assert out.exists() and out.suffix == target and out.stat().st_size > 0


# ---------- office (本地可验证部分) ----------
def test_office_xlsx_to_csv(tmp_path) -> None:
    src = _make_xlsx(tmp_path / "s.xlsx")
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(OFFICE_HANDLERS["xlsx_to_csv"], src, out_dir, ".csv")
    out = Path(res["output_path"])
    assert out.exists() and out.suffix == ".csv" and out.stat().st_size > 0
    assert "张三" in out.read_text(encoding="utf-8")


def test_office_docx_to_txt(tmp_path) -> None:
    src = _make_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(OFFICE_HANDLERS["office_to_txt"], src, out_dir, ".txt")
    out = Path(res["output_path"])
    assert "Hello 文档" in out.read_text(encoding="utf-8")


@pytest.mark.skipif(shutil.which(settings.libreoffice_bin) is None, reason="libreoffice 未安装")
def test_office_docx_to_pdf(tmp_path) -> None:
    src = _make_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"; out_dir.mkdir()
    res = _run(OFFICE_HANDLERS["office_to_pdf"], src, out_dir, ".pdf")
    out = Path(res["output_path"])
    assert out.read_bytes()[:4] == b"%PDF"


# ---------- 端到端（真实 handler 经 API）----------
def test_e2e_png_to_jpg(client, tmp_path) -> None:
    from PIL import Image
    p = tmp_path / "e.png"
    Image.new("RGB", (8, 8), (9, 9, 9)).save(str(p))
    with open(p, "rb") as f:
        r = client.post("/api/files", files={"file": ("e.png", f.read(), "image/png")})
    assert r.status_code == 200, r.text
    fid = r.json()["file_id"]
    r = client.post("/api/jobs", json={"file_id": fid, "target_ext": ".jpg"})
    jid = r.json()["job_id"]
    g = client.get(f"/api/jobs/{jid}").json()
    assert g["status"] == "succeeded"
    d = client.get(f"/api/jobs/{jid}/download")
    assert d.status_code == 200 and d.headers["content-type"].startswith("image/jpeg")
